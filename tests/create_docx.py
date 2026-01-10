from docx import Document

doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test paragraph in a DOCX file.')
doc.save(r'c:\Users\52648\Documents\02_baijiahao\AI_Resume_Sniper\tests\data\test.docx')
