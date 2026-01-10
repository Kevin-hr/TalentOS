"""
DOCX Parser / DOCX解析器

Document parser plugin for Microsoft Word files.
"""

import os
from typing import Dict, Optional
import docx

from src.interfaces.idocument_parser import IDocumentParser, ParsedDocument
from src.core.exceptions import ParseError, UnsupportedFormatError


class DOCXParser(IDocumentParser):
    """
    DOCX document parser using python-docx.

    Extracts text from Word documents (.docx).
    Note: Legacy .doc format is not supported.
    """

    PARSER_NAME = "docx"
    SUPPORTED_FORMATS = ['.docx']

    def __init__(self, extract_tables: bool = True, **kwargs):
        """
        Initialize DOCX parser.

        Args:
            extract_tables: Whether to extract table data
            **kwargs: Additional parameters
        """
        self._extract_tables = extract_tables

    @property
    def parser_name(self) -> str:
        return self.PARSER_NAME

    @property
    def supported_formats(self) -> list:
        return self.SUPPORTED_FORMATS

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a DOCX file and extract text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ParsedDocument object
        """
        if not os.path.exists(file_path):
            raise ParseError(f"File not found: {file_path}")

        # Validate file extension
        _, ext = os.path.splitext(file_path)
        if ext.lower() not in self.SUPPORTED_FORMATS:
            raise UnsupportedFormatError(f"DOCX parser only supports .docx, got: {ext}")

        try:
            doc = docx.Document(file_path)
            text = ""
            metadata = {
                "paragraphs": 0,
                "tables": 0,
                "sections": 0
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    metadata["paragraphs"] += 1

            # Extract tables
            if self._extract_tables:
                for table in doc.tables:
                    metadata["tables"] += 1
                    text += self._format_table(table) + "\n"

            return ParsedDocument(
                content=text.strip(),
                file_path=file_path,
                file_type="docx",
                metadata=metadata
            )

        except Exception as e:
            raise ParseError(f"Failed to parse DOCX {file_path}: {e}")

    def validate_file(self, file_path: str) -> bool:
        """Check if file is a valid DOCX."""
        if not os.path.exists(file_path):
            return False

        _, ext = os.path.splitext(file_path)
        return ext.lower() == '.docx'

    def parse_content(self, content: bytes, file_extension: str) -> str:
        """
        Parse DOCX content from raw bytes.

        Args:
            content: Raw DOCX bytes
            file_extension: File extension (should be '.docx')

        Returns:
            Extracted text content
        """
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            doc = self.parse(tmp_path)
            return doc.content
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    def _format_table(self, table) -> str:
        """Format extracted table data as text."""
        lines = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)

            if row_data:
                lines.append(" | ".join(row_data))

        return "\n".join(lines)
